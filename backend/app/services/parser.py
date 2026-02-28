"""
Document Parser Service
Extracts text from PDF and DOCX files
"""
import io
import tempfile
from typing import Optional
from pathlib import Path

import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class ParsedDocument:
    def __init__(
        self,
        raw_text: str,
        pages: list[dict],
        tables: list[dict],
        metadata: dict
    ):
        self.raw_text = raw_text
        self.pages = pages
        self.tables = tables
        self.metadata = metadata


class DocumentParser:
    """Parse PDF and DOCX documents to extract text content."""
    
    def __init__(self):
        self.min_text_length = 100  # Minimum text to consider valid extraction
    
    async def parse(self, file_content: bytes, file_type: str) -> ParsedDocument:
        """Parse document and extract text."""
        file_type = file_type.upper()
        
        if file_type == "PDF":
            return await self._parse_pdf(file_content)
        elif file_type in ("DOCX", "DOC"):
            return await self._parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _parse_pdf(self, content: bytes) -> ParsedDocument:
        """Parse PDF document."""
        pages = []
        tables = []
        all_text = []
        
        # Try pdfplumber first (better for text-based PDFs)
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                pages.append({
                    "page_num": i + 1,
                    "content": page_text
                })
                all_text.append(page_text)
                
                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append({
                        "page": i + 1,
                        "rows": table
                    })
        
        raw_text = "\n\n".join(all_text)
        
        # Check if we should force OCR
        has_images = await self._check_for_images(content)
        text_length = len(raw_text.strip())
        
        print(f"[PARSER] Extracted text length: {text_length}, Has images: {has_images}")
        
        # If very little text was found OR we have many images but short text, trigger OCR
        if (text_length < self.min_text_length or (has_images and text_length < 500)) and OCR_AVAILABLE:
            print("[PARSER] Text extraction poor. Triggering OCR...")
            raw_text, pages = await self._ocr_pdf(content)
            print(f"[PARSER] OCR Finished. New text length: {len(raw_text)}")
        
        metadata = {
            "page_count": len(pages),
            "has_tables": len(tables) > 0,
            "has_images": has_images,
            "method": "OCR" if len(raw_text) > text_length else "Direct"
        }
        
        return ParsedDocument(raw_text, pages, tables, metadata)
    
    async def _ocr_pdf(self, content: bytes) -> tuple[str, list[dict]]:
        """Perform OCR on scanned PDF."""
        pages = []
        all_text = []
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            
            import numpy as np
            import cv2
            
            for i, page in enumerate(doc):
                # Render page to image with higher DPI
                pix = page.get_pixmap(dpi=300)
                img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                
                # Convert to grayscale
                gray = cv2.cvtColor(img_data, cv2.COLOR_RGB2GRAY)
                
                # Apply thresholding to get black text on white background
                _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                # Optional: Denoising
                kernel = np.ones((1, 1), np.uint8)
                opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
                
                processed_img = Image.fromarray(opening)
                
                # OCR configuration
                # PSM 3: Fully automatic page segmentation, but no OSD. (Default)
                # PSM 6: Assume a single uniform block of text.
                # PSM 1: Automatic page segmentation with OSD.
                custom_config = r'--oem 3 --psm 3'
                
                # OCR
                try:
                    from app.core.config import get_settings
                    settings = get_settings()
                    if settings.tesseract_path:
                        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path
                        
                    page_text = pytesseract.image_to_string(processed_img, config=custom_config)
                except Exception as e:
                    print(f"OCR Error on page {i+1}: {e}")
                    page_text = "[OCR Failed: Tesseract binary not found or error]"
                    
                pages.append({
                    "page_num": i + 1,
                    "content": page_text
                })
                all_text.append(page_text)
            
            doc.close()
        except Exception as e:
            print(f"OCR Critical Error: {e}")
            return "", []

        return "\n\n".join(all_text), pages
    
    async def _check_for_images(self, content: bytes) -> bool:
        """Check if PDF contains images."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                if page.get_images():
                    doc.close()
                    return True
            doc.close()
        except Exception:
            pass
        return False
    
    async def _parse_docx(self, content: bytes) -> ParsedDocument:
        """Parse DOCX document and handle embedded images with OCR if needed."""
        doc = DocxDocument(io.BytesIO(content))
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
                
        # Extract tables early to include in raw_text
        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
                row_text = " | ".join([c.strip() for c in cells if c and c.strip()])
                if row_text:
                    paragraphs.append(row_text)
            tables.append({
                "index": i,
                "rows": rows
            })
            
        raw_text = "\n\n".join(paragraphs)
        text_length = len(raw_text.strip())
        
        # Extract images and perform OCR if text is poor
        ocr_text_parts = []
        image_count = 0
        
        # If text is too short, look for images inside the docx
        if text_length < self.min_text_length and OCR_AVAILABLE:
            print(f"[PARSER] DOCX text poor ({text_length} chars). Checking for images...")
            
            # Use OpenCV and Numpy for preprocessing
            import numpy as np
            import cv2
            
            # Images are hidden in document relations
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    try:
                        image_bytes = rel.target_part.blob
                        img = Image.open(io.BytesIO(image_bytes))
                        
                        # Process image for better OCR
                        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        
                        processed_img = Image.fromarray(thresh)
                        
                        # OCR configuration
                        custom_config = r'--oem 3 --psm 3'
                        
                        from app.core.config import get_settings
                        settings = get_settings()
                        if settings.tesseract_path:
                            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path
                            
                        text = pytesseract.image_to_string(processed_img, config=custom_config)
                        if text.strip():
                            ocr_text_parts.append(text)
                    except Exception as e:
                        print(f"[PARSER] Error OCRing DOCX image {image_count}: {e}")
            
            if ocr_text_parts:
                print(f"[PARSER] Extracted {len(ocr_text_parts)} texts from DOCX images.")
                raw_text = raw_text + "\n\n" + "\n\n".join(ocr_text_parts)
        
        # DOCX doesn't have reliable page breaks, treat as few chunks
        pages = [{
            "page_num": 1,
            "content": raw_text
        }]
        
        metadata = {
            "page_count": 1,
            "has_tables": len(tables) > 0,
            "paragraph_count": len(paragraphs),
            "image_count": image_count,
            "method": "OCR" if ocr_text_parts else "Direct"
        }
        
        return ParsedDocument(raw_text, pages, tables, metadata)


# Singleton instance
_parser: Optional[DocumentParser] = None


def get_parser() -> DocumentParser:
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
