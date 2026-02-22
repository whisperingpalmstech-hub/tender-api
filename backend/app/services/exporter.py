"""
Enterprise Bid Submission Export Service
Generates industry-standard DOCX proposals with dynamic company data,
professional formatting, multi-page layout, and compliance matrices.
"""
import io
import os
import re
import json
from typing import List, Dict, Optional, Any
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class CompanyProfile:
    """Company branding configuration."""
    
    def __init__(
        self,
        name: str = "TechSolutions India Pvt Ltd",
        tagline: str = "Enterprise Digital Transformation Partners",
        address: str = "Level 5, Cyber Tower, IT Park, Mumbai - 400076, Maharashtra",
        phone: str = "+91 22 4567 8900",
        email: str = "proposals@techsolutions.in",
        website: str = "www.techsolutions.in",
        logo_path: Optional[str] = None,
        primary_color: str = "#0ea5e9",  # Sky blue
        accent_color: str = "#6366f1",   # Indigo
    ):
        self.name = name
        self.tagline = tagline
        self.address = address
        self.phone = phone
        self.email = email
        self.website = website
        self.logo_path = logo_path
        self.primary_color = self._hex_to_rgb(primary_color)
        self.accent_color = self._hex_to_rgb(accent_color)
    
    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor."""
        hex_color = hex_color.lstrip('#')
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )


class ExportService:
    """Generate enterprise-grade DOCX bid submissions with dynamic data."""
    
    def __init__(self, company: Optional[CompanyProfile] = None, knowledge_base: Optional[List[Dict]] = None):
        self.company = company or CompanyProfile()
        self.kb = knowledge_base or []
        
        # Font settings
        self.font_primary = "Calibri"
        self.font_heading = "Cambria"
        
        # Size settings
        self.size_title = Pt(28)
        self.size_h1 = Pt(16)
        self.size_h2 = Pt(13)
        self.size_h3 = Pt(11)
        self.size_body = Pt(10.5)
        self.size_small = Pt(8.5)
        self.size_table = Pt(9.5)
        
        # Colors
        self.color_text = RGBColor(0x1e, 0x1e, 0x2e)
        self.color_muted = RGBColor(0x6b, 0x72, 0x80)
        self.color_light_bg = RGBColor(0xf3, 0xf4, 0xf6)
        self.color_success = RGBColor(0x05, 0x96, 0x69)
        self.color_warning = RGBColor(0xd9, 0x77, 0x06)
        self.color_white = RGBColor(0xFF, 0xFF, 0xFF)
    
    def _get_kb_items(self, category: str) -> List[Dict]:
        """Get knowledge base items by category."""
        return [item for item in self.kb if item.get('category', '').lower() == category.lower()]
    
    def _setup_styles(self, doc: Document):
        """Configure document styles for enterprise look."""
        
        # Normal style
        style = doc.styles['Normal']
        style.font.name = self.font_primary
        style.font.size = self.size_body
        style.font.color.rgb = self.color_text
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        
        # Heading 1 — Section headers
        h1 = doc.styles['Heading 1']
        h1.font.name = self.font_heading
        h1.font.size = self.size_h1
        h1.font.bold = True
        h1.font.color.rgb = self.company.primary_color
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(10)
        # Add bottom border to H1
        pPr = h1.element.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="{self._rgb_to_hex(self.company.primary_color)}"/></w:pBdr>')
        pPr.append(pBdr)
        
        # Heading 2 — Subsection
        h2 = doc.styles['Heading 2']
        h2.font.name = self.font_heading
        h2.font.size = self.size_h2
        h2.font.bold = True
        h2.font.color.rgb = self.company.accent_color
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(6)
        
        # Heading 3
        h3 = doc.styles['Heading 3']
        h3.font.name = self.font_primary
        h3.font.size = self.size_h3
        h3.font.bold = True
        h3.font.color.rgb = self.color_text
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(4)
    
    def _add_formatted_text(self, doc: Document, text: str):
        """Add text with smart formatting — handles bullet points, numbered lists, and paragraphs."""
        if not text:
            return
        lines = text.strip().split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            # Detect bullet points
            if stripped.startswith(('- ', '• ', '* ', '→ ')):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(stripped.lstrip('-•*→ ').strip())
                run.font.size = self.size_body
            # Detect numbered lists
            elif re.match(r'^\d+[\.\)]\s', stripped):
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(re.sub(r'^\d+[\.\)]\s*', '', stripped))
                run.font.size = self.size_body
            else:
                p = doc.add_paragraph(stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.runs[0].font.size = self.size_body
    
    def _add_section_divider(self, doc: Document):
        """Add a thin decorative divider line."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 50)
        run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
        run.font.size = Pt(6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    
    def _add_header_footer(self, doc: Document, tender_name: str):
        """Add professional header and footer to all sections."""
        for section in doc.sections:
            # Page margins
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            
            # Header
            header = section.header
            header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Company name in header
            cell_left = header_table.cell(0, 0)
            p = cell_left.paragraphs[0]
            run = p.add_run(self.company.name)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = self.company.primary_color
            
            # Tender name in header
            cell_right = header_table.cell(0, 1)
            p = cell_right.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(tender_name)
            run.font.size = Pt(9)
            run.font.color.rgb = self.color_muted
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = footer_para.add_run("CONFIDENTIAL | ")
            run.font.size = Pt(8)
            run.font.color.rgb = self.color_muted
            
            run = footer_para.add_run(f"© {datetime.now().year} {self.company.name}")
            run.font.size = Pt(8)
            run.font.color.rgb = self.color_muted
    
    def _add_cover_page(self, doc: Document, tender_name: str, recipient_name: str = ""):
        """Create professional cover page."""
        
        # Top spacing
        doc.add_paragraph().paragraph_format.space_before = Pt(80)
        
        # Company Logo placeholder (if logo path exists)
        if self.company.logo_path and os.path.exists(self.company.logo_path):
            try:
                # Add logo with a bit more padding
                doc.add_paragraph().paragraph_format.space_after = Pt(20)
                doc.add_picture(self.company.logo_path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph().paragraph_format.space_after = Pt(20)
            except Exception:
                pass
        
        # Company Name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company.name.upper())
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = self.company.primary_color
        run.font.name = self.font_heading
        
        # Tagline
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company.tagline)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = self.color_muted
        p.paragraph_format.space_after = Pt(60)
        
        # Decorative line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = self.company.accent_color
        p.paragraph_format.space_after = Pt(40)
        
        # Document Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tender_name)
        run.font.size = self.size_title
        run.font.bold = True
        run.font.color.rgb = self.color_text
        run.font.name = self.font_heading
        
        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Technical & Commercial Proposal")
        run.font.size = Pt(16)
        run.font.color.rgb = self.company.accent_color
        p.paragraph_format.space_after = Pt(80)
        
        # Details Table
        details_table = doc.add_table(rows=4, cols=2)
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        details = [
            ("Submitted To:", recipient_name or "Valued Client"),
            ("Submitted By:", self.company.name),
            ("Date:", datetime.now().strftime('%d %B %Y')),
            ("Version:", "1.0"),
        ]
        
        for i, (label, value) in enumerate(details):
            # Label cell
            cell_label = details_table.cell(i, 0)
            p = cell_label.paragraphs[0]
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.company.primary_color
            cell_label.width = Inches(1.5)
            
            # Value cell
            cell_value = details_table.cell(i, 1)
            p = cell_value.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(11)
            cell_value.width = Inches(3)
        
        # Contact info at bottom
        doc.add_paragraph().paragraph_format.space_before = Pt(80)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Clean up address for single line
        addr = self.company.address.replace('\n', ' | ').replace('\r', '')
        run = p.add_run(addr)
        run.font.size = Pt(9)
        run.font.color.rgb = self.color_muted
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"📞 {self.company.phone}  |  ✉ {self.company.email}  |  🌐 {self.company.website}")
        run.font.size = Pt(9)
        run.font.color.rgb = self.color_muted
        
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, sections: List[str]):
        """Add dynamic table of contents based on actual sections."""
        doc.add_heading("Table of Contents", level=1)
        
        toc_table = doc.add_table(rows=len(sections), cols=2)
        
        for i, title in enumerate(sections):
            row = toc_table.rows[i]
            
            # Section entry
            cell = row.cells[0]
            p = cell.paragraphs[0]
            run = p.add_run(f"{i+1}.  {title}")
            run.font.size = Pt(11)
            run.font.name = self.font_primary
            if i == 0:
                run.font.bold = True
            cell.width = Inches(5.5)
            p.paragraph_format.space_after = Pt(4)
            
            # Dotted leader
            cell = row.cells[1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("· · ·")
            run.font.size = Pt(10)
            run.font.color.rgb = self.color_muted
            cell.width = Inches(1)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, tender_name: str, match_summary: Optional[Dict] = None):
        """Add executive summary dynamically from KB data."""
        doc.add_heading("1. Executive Summary", level=1)
        
        # Dynamic intro using company data
        cert_items = self._get_kb_items('Certifications')
        cert_names = []
        for item in cert_items:
            title = item.get('title', '')
            if 'ISO' in title or 'CMMI' in title or 'STQC' in title:
                cert_names.append(title.split(' - ')[0].split(' Certification')[0].strip())
        
        financial_items = self._get_kb_items('Financial')
        turnover_info = ""
        for item in financial_items:
            if 'turnover' in item.get('title', '').lower():
                turnover_info = item.get('content', '')
                break
        
        summary_text = (
            f"{self.company.name} is pleased to present this comprehensive Technical & Commercial "
            f"Proposal in response to \"{tender_name}\". "
            f"As an established enterprise with deep domain expertise, we bring a proven track record "
            f"of successful project delivery across government and private sectors."
        )
        p = doc.add_paragraph(summary_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if cert_names:
            cert_text = (
                f"We are certified under {', '.join(cert_names[:3])}"
                f"{' and more' if len(cert_names) > 3 else ''}, "
                f"demonstrating our commitment to international quality and security standards."
            )
            p = doc.add_paragraph(cert_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Key Strengths
        doc.add_heading("Key Strengths", level=2)
        
        strengths = [
            "Full compliance with all eligibility, technical, and documentation requirements",
            f"{'|'.join(cert_names[:4])}" if cert_names else "Multiple ISO certifications ensuring quality delivery",
            "Experienced team of certified domain professionals",
            "Competitive pricing with transparent cost structure",
            "Proven track record of timely delivery and client satisfaction",
        ]
        if turnover_info:
            strengths.insert(2, "Strong financial standing with consistent year-over-year growth")
        
        for s in strengths:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(s.replace('|', ', '))
            run.font.size = self.size_body
        
        # Match Score Summary (if available)
        if match_summary:
            doc.add_heading("Compliance Snapshot", level=2)
            snap_table = doc.add_table(rows=1, cols=4)
            snap_table.style = 'Table Grid'
            snap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = [("Category", 1.5), ("Match %", 1), ("Status", 1), ("Remarks", 2.5)]
            for i, (h, w) in enumerate(headers):
                cell = snap_table.rows[0].cells[i]
                self._set_cell_shading(cell, self._rgb_to_hex(self.company.primary_color))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = self.color_white
                cell.width = Inches(w)
            
            cats = [
                ("Eligibility", match_summary.get('eligibility_match', 0)),
                ("Technical", match_summary.get('technical_match', 0)),
                ("Compliance", match_summary.get('compliance_match', 0)),
            ]
            for cat_name, pct in cats:
                row = snap_table.add_row().cells
                row[0].paragraphs[0].add_run(cat_name).font.size = Pt(9)
                row[1].paragraphs[0].add_run(f"{pct}%").font.size = Pt(9)
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                status = "Compliant" if pct >= 50 else "Partial"
                run = row[2].paragraphs[0].add_run(status)
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = self.color_success if pct >= 50 else self.color_warning
                row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[3].paragraphs[0].add_run("Meets stated requirements" if pct >= 50 else "Partial coverage — details in sections below").font.size = Pt(9)
            
            # Overall
            overall_row = snap_table.add_row().cells
            for cell in overall_row:
                self._set_cell_shading(cell, "EEF2FF")
            overall_row[0].paragraphs[0].add_run("Overall").font.size = Pt(9)
            overall_row[0].paragraphs[0].runs[0].font.bold = True
            overall_pct = match_summary.get('overall_match', 0)
            run = overall_row[1].paragraphs[0].add_run(f"{overall_pct}%")
            run.font.size = Pt(9)
            run.font.bold = True
            overall_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = overall_row[2].paragraphs[0].add_run("Ready" if overall_pct >= 50 else "Review")
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = self.color_success if overall_pct >= 50 else self.color_warning
            overall_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            overall_row[3].paragraphs[0].add_run("Comprehensive response submitted").font.size = Pt(9)
        
        doc.add_page_break()
    
    def _add_company_overview(self, doc: Document, company_data: Optional[Dict] = None):
        """Add dynamic company overview from KB and profile data."""
        doc.add_heading("2. Company Overview", level=1)
        
        # Pull dynamic content from KB
        legal_items = self._get_kb_items('Legal')
        registration_text = ""
        for item in legal_items:
            if 'registration' in item.get('title', '').lower() and 'legal' in item.get('title', '').lower():
                registration_text = item.get('content', '')
                break
        
        if registration_text:
            self._add_formatted_text(doc, registration_text)
        else:
            p = doc.add_paragraph(
                f"{self.company.name} is a leading enterprise solutions provider committed to "
                f"delivering excellence and driving digital transformation for clients across sectors."
            )
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Company Details Table
        doc.add_heading("Organization Details", level=2)
        
        details = [
            ("Organization Name", self.company.name),
            ("Registered Address", self.company.address.replace('\\n', ', ').replace('\n', ', ')),
            ("Contact Number", self.company.phone),
            ("Email Address", self.company.email),
            ("Website", self.company.website),
        ]
        # Add capabilities if available
        if company_data and company_data.get('capabilities'):
            caps = company_data['capabilities']
            if isinstance(caps, list):
                details.append(("Core Capabilities", ", ".join(caps[:8])))
        
        details_table = doc.add_table(rows=len(details), cols=2)
        details_table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(details):
            row = details_table.rows[i]
            
            # Label
            cell = row.cells[0]
            self._set_cell_shading(cell, "E5E7EB")
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(9.5)
            cell.width = Inches(2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Value
            cell = row.cells[1]
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(9.5)
            cell.width = Inches(4.5)
        
        # --- Certifications Sub-section ---
        cert_items = self._get_kb_items('Certifications')
        if cert_items:
            doc.add_heading("Certifications & Accreditations", level=2)
            cert_table = doc.add_table(rows=1, cols=3)
            cert_table.style = 'Table Grid'
            
            for ci, (h, w) in enumerate([("Certification", 3), ("Details", 2.5), ("Status", 1)]):
                cell = cert_table.rows[0].cells[ci]
                self._set_cell_shading(cell, self._rgb_to_hex(self.company.primary_color))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = self.color_white
                cell.width = Inches(w)
            
            for item in cert_items:
                row_cells = cert_table.add_row().cells
                row_cells[0].paragraphs[0].add_run(item.get('title', '')).font.size = Pt(9)
                content = item.get('content', '')
                # Extract key details
                summary = content[:150] + '...' if len(content) > 150 else content
                row_cells[1].paragraphs[0].add_run(summary).font.size = Pt(8.5)
                run = row_cells[2].paragraphs[0].add_run("Valid")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = self.color_success
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- Financial Strength Sub-section ---
        financial_items = self._get_kb_items('Financial')
        if financial_items:
            doc.add_heading("Financial Strength", level=2)
            for item in financial_items:
                doc.add_heading(item.get('title', ''), level=3)
                self._add_formatted_text(doc, item.get('content', ''))
        
        # --- Past Performance / Experience Sub-section ---
        experience_items = self._get_kb_items('Experience')
        if experience_items:
            doc.add_heading("Past Performance & Experience", level=2)
            for idx, item in enumerate(experience_items, 1):
                doc.add_heading(f"{idx}. {item.get('title', 'Project')}", level=3)
                self._add_formatted_text(doc, item.get('content', ''))
                if idx < len(experience_items):
                    self._add_section_divider(doc)
        
        doc.add_page_break()
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set cell background color."""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_compliance_section(
        self, 
        doc: Document, 
        section_num: int,
        section_title: str, 
        responses: List[Dict]
    ):
        """Add a compliance matrix section with formatted responses."""
        doc.add_heading(f"{section_num}. {section_title}", level=1)
        
        if not responses:
            p = doc.add_paragraph("No requirements in this category.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = self.color_muted
            return
        
        # Summary stats
        answered = sum(1 for r in responses if r.get('response_text'))
        p = doc.add_paragraph(
            f"This section addresses {len(responses)} requirements under {section_title}. "
            f"{answered} of {len(responses)} requirements have been responded to."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()
        
        # Compliance Summary Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Header row
        header_cells = table.rows[0].cells
        headers = [("S.No", 0.5), ("Requirement", 2.8), ("Response Summary", 2.7), ("Status", 0.7)]
        
        for i, (text, width) in enumerate(headers):
            cell = header_cells[i]
            self._set_cell_shading(cell, self._rgb_to_hex(self.company.primary_color))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = self.color_white
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        long_responses = []  # Track responses that need detailed pages
        
        # Data rows
        for idx, response in enumerate(responses, 1):
            req = response.get('requirement', {})
            row_cells = table.add_row().cells
            
            # S.No
            cell = row_cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(idx))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
            # Requirement
            cell = row_cells[1]
            p = cell.paragraphs[0]
            req_text = req.get('requirement_text', 'N/A')
            run = p.add_run(req_text[:200] + ('...' if len(req_text) > 200 else ''))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
            # Response — show summary in table, detail later
            cell = row_cells[2]
            p = cell.paragraphs[0]
            response_text = response.get('response_text', '')
            
            if response_text:
                # Show first 200 chars in table
                summary = response_text[:200].replace('\n', ' ')
                if len(response_text) > 200:
                    summary += '... (see details below)'
                    long_responses.append((idx, req_text, response_text))
                run = p.add_run(summary)
                run.font.size = Pt(9)
            else:
                run = p.add_run("Response pending")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = self.color_muted
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
            # Compliance status
            cell = row_cells[3]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            status = "✓" if response_text else "○"
            run = p.add_run(status)
            run.font.size = Pt(11)
            run.font.bold = True
            if response_text:
                run.font.color.rgb = self.color_success
            else:
                run.font.color.rgb = self.color_muted
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Alternate row shading
            if idx % 2 == 0:
                for cell in row_cells:
                    self._set_cell_shading(cell, "F9FAFB")
        
        # -- Detailed Responses (for long answers) --
        if long_responses:
            doc.add_paragraph()
            doc.add_heading(f"{section_title} — Detailed Responses", level=2)
            
            for idx, req_text, resp_text in long_responses:
                doc.add_heading(f"Requirement {idx}", level=3)
                
                # Show the requirement in a highlighted box
                req_p = doc.add_paragraph()
                run = req_p.add_run("Requirement: ")
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = self.company.primary_color
                run = req_p.add_run(req_text)
                run.font.size = Pt(9.5)
                run.font.italic = True
                
                # Add the full response with formatting
                resp_heading = doc.add_paragraph()
                run = resp_heading.add_run("Our Response:")
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = self.company.accent_color
                
                self._add_formatted_text(doc, resp_text)
                self._add_section_divider(doc)
        
        doc.add_paragraph()
    
    def _rgb_to_hex(self, rgb_color: RGBColor) -> str:
        """Convert RGBColor to hex string."""
        return f"{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"
    
    def _add_signature_section(self, doc: Document):
        """Add signature and authorization section."""
        doc.add_heading("Declaration & Authorization", level=1)
        
        declaration = f"""
We hereby declare that the information provided in this proposal is true and accurate to the best of our knowledge. {self.company.name} is committed to fulfilling all the requirements as stated in this document.

We understand and agree to abide by all terms and conditions of the tender. This proposal shall remain valid for a period of 90 days from the date of submission.
        """.strip()
        
        p = doc.add_paragraph(declaration)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # Signature block
        sig_table = doc.add_table(rows=4, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        sig_details = [
            ("For:", self.company.name),
            ("Name:", "______________________"),
            ("Designation:", "______________________"),
            ("Date:", datetime.now().strftime('%d %B %Y')),
        ]
        
        for i, (label, value) in enumerate(sig_details):
            row = sig_table.rows[i]
            
            cell = row.cells[0]
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(11)
            cell.width = Inches(1.2)
            
            cell = row.cells[1]
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(11)
            if i == 0:
                run.font.bold = True
            cell.width = Inches(3)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("(Authorized Signatory)")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = self.color_muted
    
    async def export_to_docx(
        self,
        tender_name: str,
        responses: List[Dict],
        requirements: List[Dict],
        company_name: str = None,
        recipient_name: str = "",
        match_summary: Optional[Dict] = None,
        company_data: Optional[Dict] = None,
    ) -> bytes:
        """Generate enterprise-grade tender response document with dynamic content."""
        
        # Update company name if provided
        if company_name:
            self.company.name = company_name
        
        doc = Document()
        
        # Setup styles
        self._setup_styles(doc)
        
        # Group responses by category (needed for TOC)
        responses_by_category = self._group_by_category(responses, requirements)
        
        # Build dynamic TOC entries
        toc_sections = ["Executive Summary", "Company Overview"]
        compliance_sections = [
            ("ELIGIBILITY", "Eligibility Criteria"),
            ("TECHNICAL", "Technical Requirements"),
            ("COMPLIANCE", "Compliance & Documentation"),
        ]
        for cat_key, cat_title in compliance_sections:
            if responses_by_category.get(cat_key):
                toc_sections.append(cat_title)
        toc_sections.append("Declaration & Authorization")
        
        # --- Build Document ---
        self._add_cover_page(doc, tender_name, recipient_name)
        self._add_header_footer(doc, tender_name)
        self._add_table_of_contents(doc, toc_sections)
        self._add_executive_summary(doc, tender_name, match_summary)
        self._add_company_overview(doc, company_data)
        
        # Add compliance sections
        section_num = 3
        for category_key, section_title in compliance_sections:
            if responses_by_category.get(category_key):
                self._add_compliance_section(
                    doc, 
                    section_num, 
                    section_title, 
                    responses_by_category[category_key]
                )
                doc.add_page_break()
                section_num += 1
        
        # Signature section
        self._add_signature_section(doc)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _group_by_category(
        self,
        responses: List[Dict],
        requirements: List[Dict]
    ) -> Dict[str, List[Dict]]:
        """Group responses by requirement category."""
        req_map = {r['id']: r for r in requirements}
        grouped = {}
        
        for response in responses:
            req = req_map.get(response.get('requirement_id'))
            category = req.get('category', 'TECHNICAL') if req else 'TECHNICAL'
            if category not in grouped:
                grouped[category] = []
            grouped[category].append({**response, 'requirement': req})
        
        return grouped


# Singleton instance
_exporter: ExportService = None


def get_exporter(company: Optional[CompanyProfile] = None, knowledge_base: Optional[List[Dict]] = None) -> ExportService:
    """Get or create exporter instance with optional KB data."""
    global _exporter
    if _exporter is None or company is not None:
        _exporter = ExportService(company, knowledge_base)
    return _exporter

