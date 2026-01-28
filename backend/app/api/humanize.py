"""
Professional AI Content Paraphraser & Humanizer
------------------------------------------------
Uses techniques from Grammarly, QuillBot, and other paraphrasing tools:
1. Synonym replacement with context awareness
2. Sentence restructuring (active/passive conversion)
3. Phrase-level paraphrasing
4. LLM-powered deep rewriting
5. Multi-pass refinement

Endpoint: POST /api/humanize
"""

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from typing import Optional, List, Tuple, Dict
import httpx
import re
import random
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
import pdfplumber
import docx
import io

from app.core.config import get_settings

settings = get_settings()
router = APIRouter(prefix="/api", tags=["humanize"])


# ============ Models ============

class HumanizeRequest(BaseModel):
    text: str = Field(..., description="Text to humanize/paraphrase")
    source_content: Optional[str] = Field(None, description="Original source for comparison")
    max_ai_percentage: float = Field(30.0, ge=0, le=100)
    max_attempts: int = Field(10, ge=1, le=20)
    style: str = Field("professional", description="professional, casual, formal, simple, academic")
    mode: str = Field("balanced", description="light, balanced, aggressive, creative")


class HumanizeResponse(BaseModel):
    success: bool
    original_text: str
    humanized_text: str
    original_ai_percentage: float
    final_ai_percentage: float
    attempts_used: int
    techniques_applied: List[str]
    message: str


# ============ SYNONYM DATABASE ============
# Context-aware synonyms (grouped by meaning/usage)

SYNONYMS = {
    # Verbs
    "use": ["employ", "apply", "utilize", "adopt", "make use of"],
    "utilize": ["use", "employ", "apply", "harness"],
    "leverage": ["use", "employ", "capitalize on", "take advantage of"],
    "implement": ["execute", "carry out", "put into practice", "deploy", "roll out"],
    "facilitate": ["enable", "help", "assist", "support", "make possible"],
    "enhance": ["improve", "boost", "strengthen", "elevate", "upgrade"],
    "optimize": ["improve", "refine", "fine-tune", "streamline"],
    "achieve": ["accomplish", "attain", "reach", "realize"],
    "ensure": ["guarantee", "make sure", "confirm", "secure"],
    "provide": ["offer", "supply", "deliver", "give", "furnish"],
    "enable": ["allow", "permit", "let", "empower", "make possible"],
    "demonstrate": ["show", "display", "illustrate", "exhibit", "prove"],
    "establish": ["set up", "create", "found", "build", "form"],
    "maintain": ["keep", "preserve", "sustain", "uphold", "retain"],
    "require": ["need", "demand", "call for", "necessitate"],
    "develop": ["create", "build", "design", "craft", "form"],
    "consider": ["think about", "examine", "look at", "review", "weigh"],
    "evaluate": ["assess", "analyze", "examine", "review", "judge"],
    "indicate": ["show", "suggest", "point to", "reveal", "signal"],
    "address": ["tackle", "handle", "deal with", "take on", "resolve"],
    
    # Adjectives
    "comprehensive": ["complete", "thorough", "full", "extensive", "all-inclusive"],
    "robust": ["strong", "solid", "sturdy", "resilient", "reliable"],
    "innovative": ["new", "novel", "creative", "original", "groundbreaking"],
    "significant": ["major", "important", "notable", "considerable", "substantial"],
    "essential": ["vital", "crucial", "key", "necessary", "fundamental"],
    "effective": ["successful", "productive", "efficient", "powerful"],
    "efficient": ["effective", "productive", "streamlined", "optimized"],
    "various": ["different", "diverse", "several", "numerous", "multiple"],
    "complex": ["complicated", "intricate", "sophisticated", "elaborate"],
    "critical": ["crucial", "vital", "key", "essential", "important"],
    "substantial": ["significant", "considerable", "major", "large", "sizable"],
    "pivotal": ["key", "crucial", "central", "vital", "essential"],
    "paramount": ["supreme", "top", "chief", "primary", "foremost"],
    "seamless": ["smooth", "effortless", "fluid", "uninterrupted"],
    "strategic": ["planned", "calculated", "deliberate", "tactical"],
    
    # Nouns  
    "methodology": ["method", "approach", "technique", "process", "system"],
    "framework": ["structure", "system", "model", "foundation"],
    "paradigm": ["model", "pattern", "example", "standard"],
    "endeavor": ["effort", "attempt", "undertaking", "venture"],
    "landscape": ["environment", "scene", "field", "arena", "sphere"],
    "synergy": ["cooperation", "teamwork", "collaboration", "partnership"],
    "implementation": ["execution", "rollout", "deployment", "application"],
    "utilization": ["use", "usage", "application", "employment"],
    
    # Adverbs
    "significantly": ["greatly", "considerably", "substantially", "noticeably"],
    "effectively": ["successfully", "well", "efficiently", "properly"],
    "subsequently": ["later", "afterward", "then", "next"],
    "predominantly": ["mainly", "mostly", "primarily", "largely"],
    "consequently": ["as a result", "therefore", "thus", "hence"],
    "additionally": ["also", "moreover", "plus", "besides", "as well"],
    "furthermore": ["also", "in addition", "plus", "what's more"],
    "however": ["but", "yet", "still", "though", "nevertheless"],
    "therefore": ["so", "thus", "hence", "as a result"],
    "moreover": ["also", "besides", "in addition", "plus"],
}

# Phrase-level paraphrasing (like QuillBot)
PHRASE_PARAPHRASES = {
    "it is important to note that": ["notably", "it's worth mentioning that", "keep in mind that", ""],
    "it should be noted that": ["note that", "importantly", "notably", ""],
    "in order to": ["to", "so as to", "for"],
    "due to the fact that": ["because", "since", "as"],
    "at this point in time": ["now", "currently", "at present"],
    "in the event that": ["if", "should", "in case"],
    "for the purpose of": ["to", "for", "in order to"],
    "with regard to": ["about", "regarding", "concerning", "on"],
    "in terms of": ["regarding", "concerning", "when it comes to", "for"],
    "a large number of": ["many", "numerous", "lots of", "plenty of"],
    "a significant amount of": ["much", "considerable", "substantial"],
    "on the other hand": ["however", "but", "conversely", "alternatively"],
    "as a result of": ["because of", "due to", "owing to", "from"],
    "in light of": ["given", "considering", "because of", "due to"],
    "with respect to": ["regarding", "about", "concerning", "for"],
    "in accordance with": ["following", "per", "as per", "according to"],
    "prior to": ["before", "ahead of", "preceding"],
    "subsequent to": ["after", "following", "post"],
    "in the process of": ["currently", "now", "busy with"],
    "take into consideration": ["consider", "think about", "account for"],
    "come to the conclusion": ["conclude", "decide", "determine"],
    "make a decision": ["decide", "choose", "determine"],
    "is able to": ["can", "is capable of"],
    "has the ability to": ["can", "is able to"],
    "there are many": ["many", "numerous", "several"],
    "it is possible that": ["possibly", "perhaps", "maybe"],
    "the fact that": ["that", "how"],
    "in spite of the fact that": ["although", "though", "despite"],
    "regardless of the fact that": ["although", "even though"],
}

# Sentence starters to vary
SENTENCE_STARTER_VARIATIONS = {
    "This": ["The", "Such", "That", ""],
    "It is": ["", "There's", "We find that"],
    "There are": ["We have", "You'll find", ""],
    "The": ["This", "Our", "Their", ""],
    "We": ["Our team", "The company", ""],
    "They": ["The team", "These professionals", ""],
}

# AI detection markers (high penalty)
AI_MARKERS = [
    r"\bdelve\b",
    r"\btapestry\b",
    r"\brealm\b",
    r"\blandscape\b",
    r"\bjourney\b",
    r"\bunlock\b",
    r"\bempower\b",
    r"\bseamless\b",
    r"\brobust\b",
    r"\bholistic\b",
    r"\bsynergy\b",
    r"\bparadigm\b",
    r"\binnovative\b",
    r"\bcutting-edge\b",
    r"\bstate-of-the-art\b",
    r"\bgame-changer\b",
    r"\bgame changing\b",
    r"\bgroundbreaking\b",
    r"\brevolutionary\b",
    r"\bunprecedented\b",
    r"\bworld-class\b",
    r"\bbest-in-class\b",
]


# ============ PARAPHRASING TECHNIQUES ============

def apply_synonym_replacement(text: str, intensity: float = 0.3) -> Tuple[str, int]:
    """
    Replace words with synonyms based on intensity (0.0-1.0).
    Higher intensity = more replacements.
    """
    words = text.split()
    replacements = 0
    result = []
    
    for word in words:
        word_lower = word.lower().strip('.,!?;:')
        
        if word_lower in SYNONYMS and random.random() < intensity:
            # Get synonyms for this word
            synonyms = SYNONYMS[word_lower]
            # Pick a random synonym
            replacement = random.choice(synonyms)
            
            # Preserve original capitalization
            if word[0].isupper():
                replacement = replacement.capitalize()
            
            # Preserve trailing punctuation
            trailing = ""
            for char in reversed(word):
                if char in '.,!?;:':
                    trailing = char + trailing
                else:
                    break
            
            result.append(replacement + trailing)
            replacements += 1
        else:
            result.append(word)
    
    return ' '.join(result), replacements


def apply_phrase_paraphrasing(text: str) -> Tuple[str, int]:
    """
    Replace common phrases with alternatives (like QuillBot).
    """
    result = text
    replacements = 0
    
    for phrase, alternatives in PHRASE_PARAPHRASES.items():
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        matches = pattern.findall(result)
        
        if matches:
            # Pick a random alternative
            replacement = random.choice(alternatives)
            result = pattern.sub(replacement, result, count=1)
            replacements += 1
    
    return result, replacements


def apply_sentence_restructuring(text: str) -> Tuple[str, int]:
    """
    Restructure sentences:
    1. Active to passive voice (sometimes)
    2. Move clauses around
    3. Split or combine sentences
    """
    sentences = re.split(r'([.!?]+)', text)
    result = []
    changes = 0
    
    i = 0
    while i < len(sentences):
        sentence = sentences[i].strip()
        punct = sentences[i + 1] if i + 1 < len(sentences) else "."
        
        if not sentence:
            i += 2
            continue
        
        # Technique: Remove redundant sentence starters
        for starter, alternatives in SENTENCE_STARTER_VARIATIONS.items():
            if sentence.startswith(starter + " ") and random.random() > 0.6:
                alt = random.choice([a for a in alternatives if a])
                if alt:
                    sentence = alt + sentence[len(starter):]
                else:
                    # Remove starter entirely if makes sense
                    rest = sentence[len(starter):].strip()
                    if rest and rest[0].islower():
                        sentence = rest[0].upper() + rest[1:]
                    else:
                        sentence = rest
                changes += 1
                break
        
        # Technique: Convert "X and Y" to "Both X and Y" or "X as well as Y"
        if " and " in sentence and random.random() > 0.7:
            if random.random() > 0.5:
                sentence = sentence.replace(" and ", " as well as ", 1)
            changes += 1
        
        # Technique: Add variety to sentence connections
        if sentence.startswith("However") and random.random() > 0.5:
            sentence = "But" + sentence[7:]
            changes += 1
        elif sentence.startswith("Therefore") and random.random() > 0.5:
            sentence = "So" + sentence[9:]
            changes += 1
        elif sentence.startswith("Additionally") and random.random() > 0.5:
            sentence = "Also" + sentence[12:]
            changes += 1
        
        result.append(sentence + punct)
        i += 2
    
    return ' '.join(result), changes


def add_contractions(text: str) -> Tuple[str, int]:
    """Add natural contractions to make text sound more human."""
    contractions = [
        (r"\bdo not\b", "don't"),
        (r"\bdoes not\b", "doesn't"),
        (r"\bcannot\b", "can't"),
        (r"\bwill not\b", "won't"),
        (r"\bshould not\b", "shouldn't"),
        (r"\bwould not\b", "wouldn't"),
        (r"\bcould not\b", "couldn't"),
        (r"\bis not\b", "isn't"),
        (r"\bare not\b", "aren't"),
        (r"\bwas not\b", "wasn't"),
        (r"\bwere not\b", "weren't"),
        (r"\bhas not\b", "hasn't"),
        (r"\bhave not\b", "haven't"),
        (r"\bhad not\b", "hadn't"),
        (r"\bit is\b", "it's"),
        (r"\bthat is\b", "that's"),
        (r"\bwhat is\b", "what's"),
        (r"\bwho is\b", "who's"),
        (r"\bthere is\b", "there's"),
        (r"\bhere is\b", "here's"),
        (r"\bthey are\b", "they're"),
        (r"\bwe are\b", "we're"),
        (r"\byou are\b", "you're"),
        (r"\bI am\b", "I'm"),
        (r"\bI have\b", "I've"),
        (r"\bI will\b", "I'll"),
        (r"\bI would\b", "I'd"),
        (r"\blet us\b", "let's"),
        (r"\bwill be\b", "will be"),  # Keep as-is
        (r"\bit will\b", "it'll"),
        (r"\bthat will\b", "that'll"),
    ]
    
    result = text
    count = 0
    
    for pattern, replacement in contractions:
        if random.random() > 0.3:  # Only apply 70% of the time for variety
            before = result
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
            if before != result:
                count += 1
    
    return result, count


def remove_ai_markers(text: str) -> Tuple[str, int]:
    """Remove words that are commonly flagged as AI-generated."""
    result = text
    count = 0
    
    # Word replacements
    ai_word_replacements = {
        "delve": "explore",
        "tapestry": "mix",
        "realm": "area",
        "landscape": "field",
        "journey": "process",
        "unlock": "discover",
        "empower": "enable",
        "seamless": "smooth",
        "robust": "strong",
        "holistic": "complete",
        "synergy": "cooperation",
        "paradigm": "approach",
        "innovative": "new",
        "cutting-edge": "modern",
        "state-of-the-art": "latest",
        "game-changer": "breakthrough",
        "groundbreaking": "major",
        "revolutionary": "significant",
        "unprecedented": "unique",
        "world-class": "excellent",
        "best-in-class": "top",
        # Additional words
        "leverage": "use",
        "utilize": "use",
        "utilization": "use",
        "facilitate": "help",
        "comprehensive": "complete",
        "dynamic": "active",
        "evolving": "growing",
        "strategic": "planned",
        "pivotal": "key",
        "paramount": "vital",
        "endeavor": "effort",
        "noteworthy": "important",
        "fortify": "strengthen",
        "fostering": "encouraging",
        "bolster": "support",
        "underscore": "show",
        "multifaceted": "varied",
        "vibrant": "lively",
        "ongoing": "current",
    }
    
    for word, replacement in ai_word_replacements.items():
        pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
        if pattern.search(result):
            result = pattern.sub(replacement, result)
            count += 1
    
    # Phrase replacements
    ai_phrase_replacements = {
        "in essence": "",
        "at its core": "",
        "plays a crucial role": "is important",
        "plays a vital role": "matters",
        "plays an important role": "helps",
        "it's worth noting": "",
        "what's more": "also",
        "in today's world": "today",
        "in the modern era": "now",
        "continues to grow": "keeps growing",
        "continues to evolve": "keeps changing",
        "further reinforced": "strengthened",
        "continually growing": "growing",
        "continually adapting": "adapting",
    }
    
    for phrase, replacement in ai_phrase_replacements.items():
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        if pattern.search(result):
            result = pattern.sub(replacement, result)
            count += 1
    
    # Clean up extra spaces
    result = re.sub(r'\s+', ' ', result).strip()
    result = re.sub(r'\s+([.,!?])', r'\1', result)
    
    return result, count


# ============ ADVANCED AI SCORE CALCULATION ============
# Inspired by GPTZero, Originality.ai, and academic AI detection papers

def calculate_ai_score(text: str) -> Tuple[float, List[str]]:
    """
    Advanced AI detection using multiple signals:
    1. Perplexity simulation (word predictability)
    2. Burstiness (sentence complexity variation)
    3. Formality score (overly polished = AI)
    4. Lexical diversity (AI tends to be repetitive)
    5. Sentence structure patterns
    6. Known AI phrases and markers
    """
    detected = []
    scores = {}  # Individual scores for each metric
    text_lower = text.lower()
    words = text.split()
    word_count = len(words)
    
    if word_count < 5:
        return 0, ["text_too_short"]
    
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip() and len(s.strip()) > 3]
    sentence_count = len(sentences)
    
    # ========== 1. BURSTINESS SCORE ==========
    # AI text has LOW burstiness (uniform complexity)
    # Human text has HIGH burstiness (varied complexity)
    burstiness_score = 0
    if sentence_count >= 2:
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_len = sum(sentence_lengths) / len(sentence_lengths)
        
        # Calculate variance
        if avg_len > 0:
            variance = sum((l - avg_len) ** 2 for l in sentence_lengths) / len(sentence_lengths)
            std_dev = variance ** 0.5
            coefficient_of_variation = std_dev / avg_len
            
            # Low variation = AI-like
            if coefficient_of_variation < 0.25:  # Very uniform
                burstiness_score = 35
                detected.append("very_uniform_sentences")
            elif coefficient_of_variation < 0.40:
                burstiness_score = 20
                detected.append("uniform_sentences")
            elif coefficient_of_variation < 0.55:
                burstiness_score = 10
                detected.append("slightly_uniform")
    
    scores['burstiness'] = burstiness_score
    
    # ========== 2. LEXICAL DIVERSITY ==========
    # AI tends to repeat words; humans use more variety
    lexical_score = 0
    unique_words = set(w.lower().strip('.,!?;:"\'-') for w in words if len(w) > 2)
    
    if word_count > 0:
        type_token_ratio = len(unique_words) / word_count
        
        # Low diversity = AI-like
        if type_token_ratio < 0.45:
            lexical_score = 25
            detected.append("low_vocabulary_diversity")
        elif type_token_ratio < 0.55:
            lexical_score = 15
            detected.append("moderate_vocabulary_diversity")
        elif type_token_ratio < 0.65:
            lexical_score = 8
    
    scores['lexical'] = lexical_score
    
    # ========== 3. FORMALITY SCORE ==========
    # AI text tends to be overly formal and polished
    formality_score = 0
    
    # Check for lack of contractions (AI often avoids them)
    contraction_patterns = [
        r"\bdon't\b", r"\bcan't\b", r"\bwon't\b", r"\bisn't\b", r"\baren't\b",
        r"\bwasn't\b", r"\bweren't\b", r"\bhasn't\b", r"\bhaven't\b", r"\bit's\b",
        r"\bthat's\b", r"\bwhat's\b", r"\bthey're\b", r"\bwe're\b", r"\bI'm\b",
        r"\bI've\b", r"\bI'll\b", r"\blet's\b", r"\bdidn't\b", r"\bcouldn't\b"
    ]
    has_contractions = any(re.search(p, text) for p in contraction_patterns)
    
    if word_count > 30 and not has_contractions:
        formality_score += 15
        detected.append("no_contractions")
    
    # Check for formal transition words (AI overuses these)
    formal_transitions = [
        (r"\bfurthermore\b", 8),
        (r"\bmoreover\b", 8),
        (r"\bnevertheless\b", 8),
        (r"\bnonetheless\b", 8),
        (r"\bconsequently\b", 7),
        (r"\bsubsequently\b", 7),
        (r"\badditionall?y\b", 5),
        (r"\bhowever\b", 3),  # Lower penalty, humans use this too
        (r"\btherefore\b", 4),
        (r"\bthus\b", 5),
        (r"\bhence\b", 6),
        (r"\bin particular\b", 4),
        (r"\bspecifically\b", 3),
        (r"\bnotably\b", 4),
        (r"\bindeed\b", 4),
        (r"\bultimately\b", 4),
        (r"\bparticularly\b", 3),
    ]
    
    for pattern, penalty in formal_transitions:
        if re.search(pattern, text_lower):
            formality_score += penalty
            detected.append(f"formal_word:{pattern[2:-2]}")
    
    scores['formality'] = min(formality_score, 40)
    
    # ========== 4. SENTENCE STARTER REPETITION ==========
    # AI often starts sentences the same way
    starter_score = 0
    if sentence_count >= 3:
        starters = []
        for s in sentences:
            words_in_sentence = s.split()
            if words_in_sentence:
                # Get first 1-2 words as starter
                starter = words_in_sentence[0].lower()
                starters.append(starter)
        
        # Count repetitions
        starter_counts = {}
        for s in starters:
            starter_counts[s] = starter_counts.get(s, 0) + 1
        
        for starter, count in starter_counts.items():
            ratio = count / len(starters)
            if ratio >= 0.5:  # Same starter for 50%+ sentences
                starter_score += 20
                detected.append(f"repetitive_starter:{starter}")
            elif ratio >= 0.33 and count >= 2:  # Same starter for 33%+ 
                starter_score += 10
                detected.append(f"common_starter:{starter}")
    
    scores['starters'] = min(starter_score, 25)
    
    # ========== 5. PERFECT STRUCTURE DETECTION ==========
    # AI text often has "perfect" paragraph structure
    structure_score = 0
    
    # Check for list-like parallel structure
    if sentence_count >= 3:
        # Check if sentences follow similar patterns
        sent_patterns = []
        for s in sentences:
            words_s = s.split()
            if len(words_s) >= 3:
                # Pattern: first word category + approximate length bucket
                pattern = f"{words_s[0].lower()}_{len(words_s)//5}"
                sent_patterns.append(pattern)
        
        if sent_patterns:
            unique_patterns = set(sent_patterns)
            if len(unique_patterns) / len(sent_patterns) < 0.5:
                structure_score += 15
                detected.append("repetitive_structure")
    
    scores['structure'] = structure_score
    
    # ========== 6. KNOWN AI PHRASES ==========
    phrase_score = 0
    
    # High-confidence AI phrases
    ai_phrases = [
        (r"it is important to note", 15),
        (r"it should be noted", 12),
        (r"it is worth mentioning", 12),
        (r"it is essential to", 8),
        (r"it is crucial to", 8),
        (r"this highlights the", 6),
        (r"this underscores", 8),
        (r"this demonstrates", 5),
        (r"in today's world", 8),
        (r"in the modern era", 8),
        (r"in conclusion", 6),
        (r"to summarize", 5),
        (r"as we have seen", 8),
        (r"as discussed", 5),
        (r"moving forward", 6),
        (r"going forward", 5),
        (r"at its core", 6),
        (r"at the heart of", 6),
        (r"plays a crucial role", 8),
        (r"plays a vital role", 8),
        (r"plays an important role", 6),
        (r"serves as a", 4),
        (r"acts as a", 3),
        (r"continues to be", 4),
        (r"remains a key", 5),
        (r"stands as", 5),
    ]
    
    for phrase, penalty in ai_phrases:
        if re.search(phrase, text_lower):
            phrase_score += penalty
            detected.append(f"ai_phrase:{phrase[:25]}")
    
    scores['phrases'] = min(phrase_score, 35)
    
    # ========== 7. KNOWN AI BUZZWORDS ==========
    buzzword_score = 0
    ai_buzzwords = [
        (r"\bdelve\b", 15),  # Very AI-specific
        (r"\btapestry\b", 12),
        (r"\blandscape\b", 5),
        (r"\bjourney\b", 4),
        (r"\bunlock\b", 5),
        (r"\bempower\b", 5),
        (r"\bseamless\b", 6),
        (r"\brobust\b", 5),
        (r"\binnovative\b", 4),
        (r"\bcomprehensive\b", 3),
        (r"\bholistic\b", 8),
        (r"\bsynergy\b", 10),
        (r"\bparadigm\b", 10),
        (r"\bcutting-edge\b", 8),
        (r"\bstate-of-the-art\b", 8),
        (r"\bgroundbreaking\b", 6),
        (r"\bunprecedented\b", 5),
        (r"\bpivotal\b", 6),
        (r"\bparamount\b", 7),
        (r"\bmultifaceted\b", 8),
        (r"\bfostering\b", 6),
        (r"\bbolster\b", 6),
        (r"\bunderscore\b", 6),
        (r"\bevolving\b", 3),
        (r"\bdynamic\b", 3),
        (r"\bstrategic\b", 3),
    ]
    
    for pattern, penalty in ai_buzzwords:
        count = len(re.findall(pattern, text_lower))
        if count > 0:
            buzzword_score += penalty * min(count, 2)
            if count > 0:
                detected.append(f"buzzword:{pattern[2:-2]}")
    
    scores['buzzwords'] = min(buzzword_score, 30)
    
    # ========== 8. COMMA DENSITY CHECK ==========
    # AI tends to use more complex sentences with more commas
    comma_score = 0
    comma_count = text.count(',')
    comma_density = comma_count / max(word_count, 1)
    
    if comma_density > 0.12:  # More than 1 comma per 8 words
        comma_score = 10
        detected.append("high_comma_density")
    elif comma_density > 0.08:
        comma_score = 5
    
    scores['commas'] = comma_score
    
    # ========== CALCULATE FINAL SCORE ==========
    total_raw = sum(scores.values())
    
    # Normalize to 0-100 scale
    # Base score plus length normalization
    if word_count < 50:
        # Short text - be more lenient
        final_score = min(total_raw * 0.8, 100)
    elif word_count < 100:
        final_score = min(total_raw * 0.9, 100)
    else:
        final_score = min(total_raw, 100)
    
    # Add debug info
    detected.append(f"[scores:{scores}]")
    
    return round(final_score, 2), detected


# ============ LLM PARAPHRASING ============

async def llm_paraphrase(text: str, style: str, mode: str, attempt: int) -> str:
    """Use LLM for intelligent paraphrasing like Grammarly/QuillBot."""
    
    if not settings.mistral_api_key:
        # No LLM - return text with basic transforms
        result = text
        result, _ = apply_synonym_replacement(result, 0.4)
        result, _ = apply_phrase_paraphrasing(result)
        result, _ = add_contractions(result)
        return result
    
    mode_instructions = {
        "light": "Make minimal changes. Only fix obvious AI patterns.",
        "balanced": "Rewrite moderately while staying close to original length.",
        "aggressive": "Completely rewrite while keeping it SHORT and simple.",
        "creative": "Rewrite naturally like a real person would speak.",
    }
    
    style_instructions = {
        "professional": "business-like but natural",
        "casual": "friendly, like texting a colleague",
        "formal": "proper but not robotic",
        "simple": "simple words, short sentences",
        "academic": "scholarly but readable",
    }
    
    original_word_count = len(text.split())
    
    prompt = f"""Rewrite this COMPLETE text to sound human-written, not AI-generated.

STRICT RULES (MUST FOLLOW):
1. Keep EXACTLY the same meaning - do not omit or summarize any information.
2. Maintain the same depth and detail as the original.
3. Use contractions where natural: don't, can't, it's, they're, we're.
4. Mix sentence lengths for a natural rhythm - some short, some longer.
5. Every single point in the original must be present in your rewrite.

WORDS TO AVOID (replace with simpler alternatives):
- leverage → use
- utilize → use  
- facilitate → help
- comprehensive → full/complete
- robust → strong
- innovative → new
- dynamic → active/changing
- strategic → planned
- evolving → growing
- seamless → smooth
- cutting-edge → modern/latest
- furthermore/moreover → also/plus
- additionally → also
- noteworthy → important
- fortify → strengthen

PHRASES TO REMOVE COMPLETELY:
- "It is important to note that"
- "It should be noted"
- "In essence"
- "plays a crucial role"
- "at its core"

STYLE: {style_instructions.get(style, style_instructions['professional'])}
MODE: {mode_instructions.get(mode, mode_instructions['balanced'])}

TEXT TO REWRITE:
{text}

REWRITTEN DOCUMENT (matching original length of approximately {original_word_count} words):"""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            # Increase tokens to ensure full response. 1 word approx 1.3 tokens.
            # We use word_count * 2 + 500 for a safe buffer.
            safe_max_tokens = max(1000, int(original_word_count * 2.5) + 500)
            
            response = await client.post(
                f"{settings.mistral_api_url}/v1/chat/completions",
                headers={"Authorization": f"Bearer {settings.mistral_api_key}"},
                json={
                    "model": settings.mistral_model,
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": safe_max_tokens,
                    "temperature": 0.6 + (attempt * 0.08),
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                paraphrased = result["choices"][0]["message"]["content"].strip()
                
                # Clean up any meta-text the LLM might add
                for prefix in ["Here's the paraphrased", "Paraphrased:", "Here is", "PARAPHRASED:", 
                               "Here's the rewritten", "Rewritten:", "REWRITTEN:", "Rewritten Document:"]:
                    if paraphrased.lower().startswith(prefix.lower()):
                        paraphrased = paraphrased[len(prefix):].strip()
                        if paraphrased.startswith(":"):
                            paraphrased = paraphrased[1:].strip()
                
                # Additional cleanup: Remove quotes if wrapped
                if paraphrased.startswith('"') and paraphrased.endswith('"'):
                    paraphrased = paraphrased[1:-1]
                
                # No longer hard-truncating to original_word_count as it cuts off sentences.
                # Instead, just return the full polished response.
                
                return paraphrased
    except Exception as e:
        print(f"[PARAPHRASE] LLM Error: {e}")
    
    # Fallback
    result = text
    result, _ = apply_synonym_replacement(result, 0.5)
    result, _ = apply_phrase_paraphrasing(result)
    result, _ = add_contractions(result)
    return result


# ============ MAIN API ENDPOINT ============

@router.post("/humanize", response_model=HumanizeResponse)
async def humanize_content(request: HumanizeRequest):
    """
    🔄 Professional Paraphrasing & AI Humanization
    
    Uses techniques from Grammarly, QuillBot, and other tools:
    - Synonym replacement
    - Phrase-level paraphrasing  
    - Sentence restructuring
    - AI marker removal
    - LLM-powered rewriting
    
    Modes:
    - light: Minimal changes
    - balanced: Moderate rewriting (default)
    - aggressive: Complete overhaul
    - creative: Conversational rewrite
    """
    
    if not request.text or len(request.text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Text must be at least 10 characters")
    
    original_text = request.text.strip()
    techniques_applied = []
    
    # Calculate original AI score
    original_ai_pct, original_patterns = calculate_ai_score(original_text)
    print(f"[HUMANIZE] Original AI score: {original_ai_pct:.1f}%")
    
    # Already good enough?
    if original_ai_pct <= request.max_ai_percentage:
        return HumanizeResponse(
            success=True,
            original_text=original_text,
            humanized_text=original_text,
            original_ai_percentage=round(original_ai_pct, 2),
            final_ai_percentage=round(original_ai_pct, 2),
            attempts_used=0,
            techniques_applied=["none_needed"],
            message=f"Text already sounds human! Score: {original_ai_pct:.1f}%"
        )
    
    # ===== STEP 1: Rule-based transformations =====
    current_text = original_text
    
    # Remove AI markers first
    current_text, n = remove_ai_markers(current_text)
    if n > 0:
        techniques_applied.append(f"ai_marker_removal:{n}")
    
    # Phrase paraphrasing
    current_text, n = apply_phrase_paraphrasing(current_text)
    if n > 0:
        techniques_applied.append(f"phrase_paraphrase:{n}")
    
    # Synonym replacement based on mode
    intensity = {"light": 0.2, "balanced": 0.35, "aggressive": 0.5, "creative": 0.45}.get(request.mode, 0.35)
    current_text, n = apply_synonym_replacement(current_text, intensity)
    if n > 0:
        techniques_applied.append(f"synonym_replace:{n}")
    
    # Sentence restructuring
    current_text, n = apply_sentence_restructuring(current_text)
    if n > 0:
        techniques_applied.append(f"restructure:{n}")
    
    # Add contractions
    current_text, n = add_contractions(current_text)
    if n > 0:
        techniques_applied.append(f"contractions:{n}")
    
    # Clean up
    current_text = re.sub(r'\s+', ' ', current_text).strip()
    current_text = re.sub(r'\s+([.,!?])', r'\1', current_text)
    
    # Check score after rule-based transforms
    current_ai_pct, _ = calculate_ai_score(current_text)
    print(f"[HUMANIZE] After rules: {current_ai_pct:.1f}%")
    
    best_text = current_text
    best_ai_pct = current_ai_pct
    attempts_used = 0
    
    # ===== STEP 2: LLM paraphrasing if needed =====
    if current_ai_pct > request.max_ai_percentage:
        techniques_applied.append("llm_paraphrase")
        
        for attempt in range(request.max_attempts):
            attempts_used = attempt + 1
            print(f"[HUMANIZE] LLM attempt {attempts_used}...")
            
            # Get LLM paraphrase
            paraphrased = await llm_paraphrase(best_text, request.style, request.mode, attempt)
            
            # Apply rule-based cleanup to LLM output
            paraphrased, _ = remove_ai_markers(paraphrased)
            paraphrased, _ = apply_phrase_paraphrasing(paraphrased)
            paraphrased, _ = add_contractions(paraphrased)
            paraphrased = re.sub(r'\s+', ' ', paraphrased).strip()
            
            # Score
            new_ai_pct, _ = calculate_ai_score(paraphrased)
            print(f"[HUMANIZE] Attempt {attempts_used}: {new_ai_pct:.1f}%")
            
            # Keep if better
            if new_ai_pct < best_ai_pct:
                best_text = paraphrased
                best_ai_pct = new_ai_pct
            
            # Success?
            if best_ai_pct <= request.max_ai_percentage:
                break
    
    success = best_ai_pct <= request.max_ai_percentage
    
    return HumanizeResponse(
        success=success,
        original_text=original_text,
        humanized_text=best_text,
        original_ai_percentage=round(original_ai_pct, 2),
        final_ai_percentage=round(best_ai_pct, 2),
        attempts_used=attempts_used,
        techniques_applied=techniques_applied,
        message=f"{'✅ Success!' if success else '⚠️ Partial:'} AI score: {original_ai_pct:.1f}% → {best_ai_pct:.1f}%"
    )


# ============ UTILITY ENDPOINTS ============

@router.get("/humanize/health")
async def humanize_health():
    """Check API status."""
    return {
        "status": "healthy",
        "llm_available": bool(settings.mistral_api_key),
        "synonyms_loaded": len(SYNONYMS),
        "phrases_loaded": len(PHRASE_PARAPHRASES),
        "modes": ["light", "balanced", "aggressive", "creative"],
        "styles": ["professional", "casual", "formal", "simple", "academic"]
    }


@router.get("/humanize/test")
async def humanize_test():
    """Test with sample AI text."""
    sample = "It is important to note that our comprehensive solution leverages cutting-edge technology to facilitate seamless integration. Furthermore, we utilize robust methodologies to implement innovative strategies."
    
    original_score, patterns = calculate_ai_score(sample)
    
    # Apply transforms
    result = sample
    result, _ = remove_ai_markers(result)
    result, _ = apply_phrase_paraphrasing(result)
    result, _ = apply_synonym_replacement(result, 0.4)
    result, _ = add_contractions(result)
    result = re.sub(r'\s+', ' ', result).strip()
    
    new_score, _ = calculate_ai_score(result)
    
    return {
        "original": sample,
        "original_score": round(original_score, 1),
        "patterns_detected": patterns,
        "transformed": result,
        "new_score": round(new_score, 1),
        "reduction": f"{original_score - new_score:.1f}%"
    }


@router.post("/humanize/file", response_model=HumanizeResponse)
async def humanize_file(
    file: UploadFile = File(...),
    max_ai_percentage: float = Form(30.0),
    max_attempts: int = Form(10),
    style: str = Form("professional"),
    mode: str = Form("balanced")
):
    """
    Handle file uploads (PDF, DOCX, TXT) for humanization.
    Extracts text and forwards to the main humanize logic.
    """
    content = ""
    filename = file.filename.lower()
    file_bytes = await file.read()
    
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                content = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".txt"):
             content = file_bytes.decode("utf-8")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT.")
    except Exception as e:
        print(f"Error parsing file: {e}")
        # raise HTTPException(status_code=400, detail=f"Error parsing file: {str(e)}")
        # For robustness, try to continue even if extraction is imperfect
        raise HTTPException(status_code=400, detail="Could not parse file content.")

    content = content.replace('\x00', '') # Clean null bytes
    if not content.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")
        
    # Limit to 20k chars as per requirement
    if len(content) > 20000:
        content = content[:20000]
        
    req = HumanizeRequest(
        text=content,
        max_ai_percentage=max_ai_percentage,
        max_attempts=max_attempts,
        style=style,
        mode=mode
    )
    
    return await humanize_content(req)


# ============ UNIFIED ENDPOINT ============

root_router = APIRouter()

@root_router.post("/humanizer")
async def humanize_unified(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    max_ai_percentage: float = Form(30.0),
    max_attempts: int = Form(10),
    style: str = Form("professional"),
    mode: str = Form("balanced")
):
    """
    Unified endpoint for Humanizer.
    Accepts EITHER a file (PDF, DOCX, TXT) OR raw text (copy-paste).
    """
    print(f"Humanize Request: file={file.filename if file else 'None'}, text_len={len(text) if text else 0}")
    content = ""
    
    # 1. Handle File
    if file:
        filename = file.filename.lower()
        print(f"Processing file: {filename}")
        try:
            file_bytes = await file.read()
            
            if filename.endswith(".pdf"):
                try:
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        content = "\n".join([page.extract_text() or "" for page in pdf.pages])
                except Exception as e:
                    print(f"PDF Error: {e}")
                    raise HTTPException(status_code=400, detail="Failed to read PDF file. It might be corrupted or password protected.")
            
            elif filename.endswith(".docx") or filename.endswith(".doc"):
                try:
                    # Note: python-docx strictly supports .docx (OOXML)
                    doc = docx.Document(io.BytesIO(file_bytes))
                    content = "\n".join([p.text for p in doc.paragraphs])
                except Exception as e:
                    print(f"Word Error for {filename}: {e}")
                    if filename.endswith(".doc"):
                        raise HTTPException(status_code=400, detail="Legacy .doc files are not supported. Please save your file as .docx (Word Document) and try again.")
                    else:
                        raise HTTPException(status_code=400, detail="Failed to read Word file. Please ensure it is a valid .docx file.")
            
            elif filename.endswith(".txt"):
                 content = file_bytes.decode("utf-8", errors='ignore')
            
            else:
                print(f"Unsupported format requested: {filename}")
                raise HTTPException(status_code=400, detail=f"Unsupported file: {filename}. Please use PDF, DOCX, or TXT.")
                
        except HTTPException:
            raise
        except Exception as e:
            print(f"General File Error: {e}")
            raise HTTPException(status_code=400, detail=f"Error parsing file: {str(e)}")
            
    # 2. Handle Text (Override file if provided, or if file empty)
    if text and text.strip():
        content = text
        
    if not content or not content.strip():
        print("No content extracted")
        raise HTTPException(status_code=400, detail="No content provided. Please upload a valid file or paste text.")
        
    # Limit check
    if len(content) > 50000:
        content = content[:50000]
        
    req = HumanizeRequest(
        text=content,
        max_ai_percentage=max_ai_percentage,
        max_attempts=max_attempts,
        style=style,
        mode=mode
    )
    
    # Return mapping to ensure frontend compatibility
    # humanize_content returns a dict (step 327 analysis)
    result_data = await humanize_content(req)
    
    # Map to frontend expectations just in case
    # If result_data is a Pydantic model (HumanizeResponse), convert to dict
    if hasattr(result_data, 'dict'):
        result_data = result_data.dict()
    
    # Ensure keys exist for the updated api.ts
    return {
        "transformed": result_data.get("transformed") or result_data.get("humanized_text") or "",
        "original_score": result_data.get("original_score") or result_data.get("original_ai_percentage") or 0,
        "new_score": result_data.get("new_score") or result_data.get("final_ai_percentage") or 0,
        "reduction": result_data.get("reduction") or "0%"
    }

